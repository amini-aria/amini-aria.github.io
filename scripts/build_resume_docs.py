# -*- coding: utf-8 -*-
"""
Builds Aria-CV-En.docx / Aria-CV-Fa.docx from scripts/resume_data.py, then
converts each to a matching .pdf with LibreOffice headless — so the PDF is
literally an export of the Word file, guaranteeing the two are identical.

Design goals (per spec): plain, non-graphical, black-and-white, compact,
small font, classic and highly readable, correct/principled use of real
font weights (not synthetic bold) and italic, ATS-safe (no tables, no
icons, no colored chips, no text boxes) — built to be parsed cleanly by
applicant-tracking systems as well as read comfortably by a human.

Fonts:
  - English: Times New Roman (classic, ATS-safe, universally available).
  - Persian: the Dana family, using its real weight files (Light /
    Regular / Medium / DemiBold / Bold / Black — see assets/fonts/) rather
    than synthetic bold, for a genuinely modern/minimalist weight
    hierarchy. Requires the Dana .ttf files to be installed wherever this
    script runs (see assets/fonts/ and the GitHub Action) so LibreOffice
    can find them when converting to PDF; on the end user's own machine,
    Word will use real Dana if installed there, or substitute a fallback
    otherwise (the underlying text/content is unaffected either way).

Usage:
    python3 build_resume_docs.py

Output:
    assets/files/Aria-CV-En.docx / .pdf
    assets/files/Aria-CV-Fa.docx / .pdf
"""

import subprocess
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, Cm, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
from resume_data import RESUME_EN, RESUME_FA  # noqa: E402
from jalali import today_jalali_string  # noqa: E402

OUT_DIR = Path(__file__).parent.parent / "assets" / "files"

INK = RGBColor(0x1A, 0x1A, 0x1A)
SECONDARY = RGBColor(0x4D, 0x4D, 0x4D)
ACCENT_BLUE = RGBColor(0x1F, 0x4E, 0x79)

# Four sizes only, matching academic-CV convention (name / section title /
# body / small print) instead of a dozen near-duplicate increments — every
# other visual distinction (primary vs. secondary content) comes from
# color and weight/italic, not size.
SIZE_NAME = 16
SIZE_SECTION = 11.5
SIZE_BODY = 10.5
SIZE_SMALL = 8

FONT_EN = "Times New Roman"

# Dana ships as separate files per weight, each declaring its OWN family
# name (confirmed via each .ttf's internal name table) rather than a single
# variable-weight family — so each weight is referenced as if it were its
# own distinct font name.
DANA = {
    "light": "Dana Light",
    "regular": "Dana",
    "medium": "Dana Medium",
    "demibold": "Dana DemiBold",
    "bold": "Dana Bold",
    "black": "Dana Black",
}


def set_rtl_paragraph(paragraph):
    """Mark a paragraph as bidi (right-to-left) at the XML level — this is
    what makes Word lay it out RTL regardless of which font is applied."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rPr.append(OxmlElement("w:rtl"))


def add_right_tab(paragraph, usable_width_cm):
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(usable_width_cm), WD_TAB_ALIGNMENT.RIGHT)


def condense(run, twips=-3):
    """Slightly tightens character spacing (negative w:spacing, in
    twentieths of a point) — used on date/period text so ranges like
    'Feb 2025 - Present' or '\u0628\u0647\u0645\u0646 \u06f1\u06f4\u06f0\u06f3 \u2013 \u0627\u06a9\u0646\u0648\u0646' read as a
    tighter, more tabular block instead of loosely spaced-out text."""
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(twips))
    rPr.append(spacing)


def add_hyperlink(paragraph, text, url, font, size, color=INK, weight_bold=False, rtl=False):
    """Inserts a real, clickable OOXML hyperlink run with NO underline and
    NO color change (matches the surrounding plain text exactly) — direct
    character formatting is applied so it overrides Word's default
    'Hyperlink' style rather than inheriting its blue-underline look."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rPr.append(rFonts)

    if weight_bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    rPr.append(color_el)
    # explicitly no underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    rPr.append(u)
    if rtl:
        rPr.append(OxmlElement("w:rtl"))

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


class ResumeBuilder:
    """One instance per language. Call .build() then .save(path)."""

    def __init__(self, data, is_fa):
        self.data = data
        self.is_fa = is_fa
        self.font = DANA["regular"] if is_fa else FONT_EN
        self.doc = Document()
        self._setup_page()

    # ---- low-level helpers -------------------------------------------------

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)
        section.footer_distance = Cm(0.5)
        self.usable_width_cm = 21.0 - 1.4 - 1.4

        normal = self.doc.styles["Normal"]
        normal.font.name = self.font
        normal.font.size = Pt(SIZE_BODY)
        normal.font.color.rgb = INK
        rpr = normal.element.get_or_add_rPr()
        szCs = rpr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rpr.append(szCs)
        szCs.set(qn("w:val"), str(int(round(SIZE_BODY * 2))))
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        rFonts.set(qn("w:cs"), self.font)
        rFonts.set(qn("w:ascii"), self.font)
        rFonts.set(qn("w:hAnsi"), self.font)

    def _para(self, align=None, space_before=0, space_after=1, rtl=None):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.0
        if rtl is None:
            rtl = self.is_fa
        if rtl:
            set_rtl_paragraph(p)
        return p

    def _font_for(self, weight):
        """weight is ignored for English (Times New Roman handles bold via
        the normal w:b flag); for Persian it picks the exact Dana file."""
        if not self.is_fa:
            return self.font
        return DANA.get(weight, DANA["regular"])

    def _run(self, paragraph, text, weight="regular", italic=False, size=SIZE_BODY, color=INK, condensed=False, force_ltr=False):
        if self.is_fa:
            # Dana has no en-dash glyph; without this, LibreOffice silently
            # falls back to DejaVu Sans for just that character, producing
            # a visibly mixed font mid-run (e.g. inside date ranges).
            text = text.replace("–", "-")
        run = paragraph.add_run(text)
        font_name = self._font_for(weight)
        run.font.name = font_name
        run.font.size = Pt(size)
        # for English, "bold"/"demibold"/"black" all just mean w:b (Times
        # New Roman has one real bold, no extra weight steps); for Persian
        # the real weight comes from font_name itself, so bold stays off
        # to avoid double-emboldening a file that's already e.g. DemiBold
        run.font.bold = (weight in ("bold", "demibold", "black")) if not self.is_fa else False
        run.font.italic = italic
        run.font.color.rgb = color
        if self.is_fa and not force_ltr:
            set_rtl_run(run)
        rpr = run._r.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        rFonts.set(qn("w:cs"), font_name)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        # CRITICAL: Word/LibreOffice size complex-script runs (Arabic and
        # Persian script) using w:szCs, NOT w:sz — python-docx's
        # run.font.size only ever sets w:sz. Without this, every Persian
        # run rendered at some default complex-script size regardless of
        # what size we asked for, which is why section titles, the footer
        # date, and mixed Latin/Persian lines all looked inconsistently
        # sized before this fix.
        half_points = str(int(round(size * 2)))
        szCs = rpr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rpr.append(szCs)
        szCs.set(qn("w:val"), half_points)
        if condensed:
            condense(run)
        return run

    def _section_title(self, text):
        p = self._para(space_before=3, space_after=1)
        self._run(p, text.upper() if not self.is_fa else text, weight="demibold", size=SIZE_SECTION, color=ACCENT_BLUE)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "5")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "333333")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(3)
        return p

    def _entry_head_line(self, left_text, right_text, left_weight, left_italic=False, left_color=INK, kicker=None):
        """Classic resume pattern: left-aligned text, right-aligned date,
        using a right tab stop rather than a table (safer for ATS parsers
        and screen readers alike). The date itself is set slightly
        condensed for a tighter, more tabular look. left_color lets a
        caller mark the left text itself as secondary (e.g. a role line
        under an org heading that's already the primary line). kicker, if
        given, is a small secondary-colored label (e.g. "Paper:") printed
        just before left_text, to disambiguate what kind of title follows."""
        p = self._para(space_before=0, space_after=0)
        add_right_tab(p, self.usable_width_cm)
        if kicker:
            self._run(p, kicker + "  ", weight="demibold", size=SIZE_SMALL, color=SECONDARY)
        self._run(p, left_text, weight=left_weight, italic=left_italic, size=SIZE_BODY, color=left_color)
        if right_text:
            self._run(p, "\t" + right_text, weight="light", italic=False, size=SIZE_BODY, color=SECONDARY, condensed=True)
        return p

    def _bullet(self, text):
        p = self._para(space_before=0, space_after=0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # prepending in logical order puts the dash at the start (right
        # side) for RTL Persian paragraphs too, unlike a trailing suffix
        # which would visually land on the wrong side
        marker = "\u2013 "
        self._run(p, marker + text, weight="light", size=SIZE_BODY)

    # ---- sections ------------------------------------------------------

    def header(self):
        d = self.data
        p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, space_before=0)
        self._run(p, d["name"], weight="black", size=SIZE_NAME, color=ACCENT_BLUE)

        p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
        self._run(p, d["role"], weight="medium", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)

        p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
        parts = d["contact_parts"]
        for i, part in enumerate(parts):
            if i > 0:
                self._run(p, "   |   ", weight="light", size=SIZE_BODY, color=SECONDARY)
            if part.get("url"):
                add_hyperlink(
                    p, part["text"], part["url"],
                    font=self._font_for("light"), size=SIZE_BODY, color=SECONDARY,
                    weight_bold=False, rtl=self.is_fa,
                )
            else:
                self._run(p, part["text"], weight="light", size=SIZE_BODY, color=SECONDARY)

    def about(self):
        self._section_title("About" if not self.is_fa else "\u062f\u0631\u0628\u0627\u0631\u0647\u200c\u0645\u0646")
        p = self._para(space_after=3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._run(p, self.data["about"], weight="regular", size=SIZE_BODY)

    def interests(self):
        self._section_title("Interests" if not self.is_fa else "\u0639\u0644\u0627\u0642\u0647\u200c\u0645\u0646\u062f\u06cc\u200c\u0647\u0627")
        p = self._para(space_after=3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._run(p, self.data["interests"], weight="regular", size=SIZE_BODY)

    def education(self):
        self._section_title("Education" if not self.is_fa else "\u062a\u062d\u0635\u06cc\u0644\u0627\u062a")
        for e in self.data["education"]:
            self._entry_head_line(e["degree"], e["period"], left_weight="demibold")
            p = self._para(space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, e["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            for note in e["notes"]:
                self._bullet(note)

    def experience(self):
        title = "Professional Experience" if not self.is_fa else "\u062a\u062c\u0631\u0628\u0647 \u062d\u0631\u0641\u0647\u200c\u0627\u06cc"
        self._section_title(title)
        for e in self.data["experience"]:
            # org is the heading (demibold/medium weight); role + period
            # follow beneath it. An entry may list several roles under the
            # same org heading (e.g. two positions held at one university)
            # instead of a single role/period pair.
            p = self._para(space_before=3, space_after=0)
            self._run(p, e["org"], weight="demibold", size=SIZE_BODY)
            if "roles" in e:
                for r in e["roles"]:
                    self._entry_head_line(r["role"], r["period"], left_weight="light", left_italic=not self.is_fa, left_color=SECONDARY)
            else:
                self._entry_head_line(e["role"], e["period"], left_weight="light", left_italic=not self.is_fa, left_color=SECONDARY)

    def research(self):
        title = "Research Experience" if not self.is_fa else "\u062a\u062c\u0631\u0628\u0647 \u067e\u0698\u0648\u0647\u0634\u06cc"
        self._section_title(title)
        for r in self.data["research"]:
            self._entry_head_line(r["title"], r["period"], left_weight="demibold")
            p = self._para(space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, r["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            if r.get("note"):
                p2 = self._para(space_after=2)
                self._run(p2, r["note"], weight="light", size=SIZE_SMALL, color=SECONDARY)

    def books(self):
        self._section_title("Books" if not self.is_fa else "\u06a9\u062a\u0627\u0628\u200c\u0647\u0627")
        default_isbn_label = "ISBN: " if not self.is_fa else "\u0634\u0627\u0628\u06a9: "
        for b in self.data["books"]:
            p = self._para(space_before=2, space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, b["citation"], weight="demibold", size=SIZE_BODY)
            p2 = self._para(space_after=2)
            self._run(p2, b["role"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            if b.get("isbn"):
                self._run(p2, "   \u00b7   ", weight="light", size=SIZE_SMALL, color=SECONDARY)
                isbn_label = b.get("isbn_label", default_isbn_label)
                self._run(p2, isbn_label, weight="light", size=SIZE_SMALL, color=SECONDARY)
                # the ISBN's digit groups must read left-to-right even inside a
                # right-to-left Persian paragraph, so this run is never RTL-flagged
                self._run(p2, b["isbn"], weight="light", size=SIZE_SMALL, color=SECONDARY, force_ltr=True)

    def conferences(self):
        self._section_title("Conferences" if not self.is_fa else "کنفرانس‌ها")
        id_label = "National Sci-Doc ID: " if not self.is_fa else "شناسه ملی سند علمی: "
        link_label = "View on Civilica" if not self.is_fa else "مشاهده در سیویلیکا"
        kicker = "Paper:" if not self.is_fa else "مقاله:"
        for c in self.data["conferences"]:
            self._entry_head_line(c["title"], c["period"], left_weight="demibold", kicker=kicker)
            p = self._para(space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, c["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            p2 = self._para(space_after=0)
            self._run(p2, c["tag"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            p3 = self._para(space_after=2)
            self._run(p3, c["note"], weight="light", size=SIZE_SMALL, color=SECONDARY)
            if c.get("doc_id"):
                self._run(p3, "  ·  " + id_label, weight="light", size=SIZE_SMALL, color=SECONDARY)
                self._run(p3, c["doc_id"], weight="light", size=SIZE_SMALL, color=SECONDARY, force_ltr=True)
            if c.get("url"):
                self._run(p3, "  ·  ", weight="light", size=SIZE_SMALL, color=SECONDARY)
                add_hyperlink(
                    p3, link_label, c["url"],
                    font=self._font_for("light"), size=SIZE_SMALL, color=ACCENT_BLUE,
                    weight_bold=False, rtl=self.is_fa,
                )

    def patents(self):
        self._section_title("Patents" if not self.is_fa else "اختراعات")
        for pt in self.data["patents"]:
            self._entry_head_line(pt["title"], pt["status"], left_weight="demibold")
            p = self._para(space_after=2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, pt["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)

    def teaching(self):
        self._section_title("Teaching Activities" if not self.is_fa else "فعالیت‌های تدریس")
        for t in self.data["teaching"]:
            self._entry_head_line(t["title"], t["period"], left_weight="demibold")
            p = self._para(space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, t["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            p2 = self._para(space_after=2)
            self._run(p2, t["note"], weight="light", size=SIZE_SMALL, color=SECONDARY)

    def honors(self):
        self._section_title("Honors & Awards" if not self.is_fa else "افتخارات و جوایز")
        for h in self.data["honors"]:
            self._entry_head_line(h["title"], h["period"], left_weight="demibold")
            p = self._para(space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, h["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            if h.get("note"):
                p2 = self._para(space_after=2)
                self._run(p2, h["note"], weight="light", size=SIZE_SMALL, color=SECONDARY)

    def _chip_group(self, groups):
        for label, items in groups.items():
            p = self._para(space_before=1, space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, label + ":  ", weight="demibold", size=SIZE_BODY)
            self._run(p, "  \u2022  ".join(items), weight="regular", size=SIZE_BODY)

    def skills(self):
        title = "Technical & Specialized Skills" if not self.is_fa else "\u0645\u0647\u0627\u0631\u062a\u200c\u0647\u0627\u06cc \u0641\u0646\u06cc \u0648 \u062a\u062e\u0635\u0635\u06cc"
        self._section_title(title)
        self._chip_group(self.data["skills"])

    def languages(self):
        self._section_title("Languages" if not self.is_fa else "\u0632\u0628\u0627\u0646\u200c\u0647\u0627")
        p = self._para(space_after=1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for i, l in enumerate(self.data["languages"]):
            if i > 0:
                self._run(p, "   |   ", weight="light", size=SIZE_BODY)
            self._run(p, l["name"], weight="demibold", size=SIZE_BODY)
            self._run(p, f' \u2013 {l["level"]}', weight="regular", size=SIZE_BODY)
        for l in self.data["languages"]:
            for t in l.get("tests", []):
                p2 = self._para(space_after=0)
                p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._run(p2, t["name"] + "  ", weight="demibold", size=SIZE_SMALL, color=SECONDARY)
                self._run(p2, t["date_range"] + "  \u00b7  ", weight="light", size=SIZE_SMALL, color=SECONDARY)
                self._run(p2, t["summary"], weight="light", size=SIZE_SMALL, color=SECONDARY)
        self.doc.paragraphs[-1].paragraph_format.space_after = Pt(3)

    def memberships(self):
        title = "Professional Memberships" if not self.is_fa else "\u0639\u0636\u0648\u06cc\u062a\u200c\u0647\u0627\u06cc \u062d\u0631\u0641\u0647\u200c\u0627\u06cc"
        self._section_title(title)
        for m in self.data["memberships"]:
            self._entry_head_line(m["role"], m["period"], left_weight="demibold")
            p = self._para(space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, m["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)

    def volunteer(self):
        title = "Voluntary & Social Activities" if not self.is_fa else "\u0633\u0648\u0627\u0628\u0642 \u062f\u0627\u0648\u0637\u0644\u0628\u0627\u0646\u0647 \u0648 \u0641\u0639\u0627\u0644\u06cc\u062a\u200c\u0647\u0627\u06cc \u0627\u062c\u062a\u0645\u0627\u0639\u06cc"
        self._section_title(title)
        for v in self.data["volunteer"]:
            self._entry_head_line(v["role"], v["period"], left_weight="demibold")
            p = self._para(space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._run(p, v["org"], weight="light", italic=not self.is_fa, size=SIZE_BODY, color=SECONDARY)
            notes = v.get("notes", [])
            for i, note in enumerate(notes):
                p2 = self._para(space_after=2 if i == len(notes) - 1 else 0)
                self._run(p2, note, weight="light", size=SIZE_SMALL, color=SECONDARY)

    def certifications(self):
        title = "Courses, Certificates & Licenses" if not self.is_fa else "دوره‌ها، گواهی‌ها و مجوزها"
        self._section_title(title)
        view_cert = "View certificate" if not self.is_fa else "مشاهده گواهی"
        for spec in self.data["certifications"]["specializations"]:
            p = self._para(space_before=2, space_after=0)
            self._run(p, spec["title"], weight="demibold", size=SIZE_BODY)
            p2 = self._para(space_after=1)
            self._run(p2, spec["meta"], weight="light", size=SIZE_SMALL, color=SECONDARY)
            for c in spec["courses"]:
                p3 = self._para(space_after=0)
                p3.paragraph_format.left_indent = Cm(0.5)
                self._run(p3, "– " + c["title"] + "   ", weight="light", size=SIZE_SMALL)
                self._run(p3, c["period"] + " · ", weight="light", size=SIZE_SMALL, color=SECONDARY)
                self._run(p3, c["id"], weight="light", size=SIZE_SMALL, color=SECONDARY, force_ltr=True)
                self._run(p3, "  ·  ", weight="light", size=SIZE_SMALL, color=SECONDARY)
                add_hyperlink(
                    p3, view_cert, c["url"],
                    font=self._font_for("light"), size=SIZE_SMALL, color=ACCENT_BLUE,
                    weight_bold=False, rtl=self.is_fa,
                )
            p4 = self._para(space_before=1, space_after=3)
            p4.paragraph_format.left_indent = Cm(0.5)
            add_hyperlink(
                p4, spec["cred_label"], spec["cred_url"],
                font=self._font_for("light"), size=SIZE_SMALL, color=ACCENT_BLUE,
                weight_bold=False, rtl=self.is_fa,
            )
        for c in self.data["certifications"]["courses"]:
            p = self._para(space_after=0)
            self._run(p, c["title"] + "   ", weight="regular", size=SIZE_SMALL)
            self._run(p, c["org"] + "  ·  ", weight="light", size=SIZE_SMALL, color=SECONDARY)
            self._run(p, c["period"] + " · ", weight="light", size=SIZE_SMALL, color=SECONDARY)
            self._run(p, c["id"], weight="light", size=SIZE_SMALL, color=SECONDARY, force_ltr=True)
            self._run(p, "  ·  ", weight="light", size=SIZE_SMALL, color=SECONDARY)
            add_hyperlink(
                p, view_cert, c["url"],
                font=self._font_for("light"), size=SIZE_SMALL, color=ACCENT_BLUE,
                weight_bold=False, rtl=self.is_fa,
            )

    def footer(self):
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if self.is_fa:
            set_rtl_paragraph(p)
            label = "\u0622\u062e\u0631\u06cc\u0646 \u0648\u06cc\u0631\u0627\u06cc\u0634: "
            stamp = today_jalali_string()
        else:
            label = "Last updated: "
            stamp = date.today().strftime("%B %-d, %Y") if sys.platform != "win32" else date.today().strftime("%B %#d, %Y")
        self._run(p, label + stamp, weight="light", size=SIZE_SMALL, color=SECONDARY, italic=False)

    def build(self):
        self.header()
        self.about()
        self.interests()
        self.education()
        self.experience()
        self.research()
        self.books()
        self.conferences()
        self.patents()
        self.teaching()
        self.honors()
        self.languages()
        self.memberships()
        self.volunteer()
        self.skills()
        self.certifications()
        self.footer()
        return self.doc

    def save(self, path):
        self.doc.save(str(path))


def convert_to_pdf(docx_path, out_dir):
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    jobs = [
        (RESUME_EN, False, "Aria-CV-En"),
        (RESUME_FA, True, "Aria-CV-Fa"),
    ]

    for data, is_fa, stem in jobs:
        builder = ResumeBuilder(data, is_fa)
        builder.build()
        docx_path = OUT_DIR / f"{stem}.docx"
        builder.save(docx_path)
        print(f"wrote {docx_path}")
        convert_to_pdf(docx_path, OUT_DIR)
        print(f"converted {stem}.docx -> {stem}.pdf")


if __name__ == "__main__":
    main()
